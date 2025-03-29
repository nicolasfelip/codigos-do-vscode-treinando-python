from docx import Document

# Criação do documento
doc = Document()

# Título do Trabalho
doc.add_heading('Resolução de Sistemas de Equações do 1º Grau', 0)

# Introdução
doc.add_paragraph(
    'Podemos resolver um sistema de equações do 1º grau com duas incógnitas utilizando dois métodos principais: '
    'Método da Substituição e Método da Adição (ou Eliminação).'
)

# Método da Substituição
doc.add_heading('Método da Substituição', level=1)
doc.add_paragraph(
    'O Método da Substituição consiste em escolher uma das equações e isolar uma das incógnitas, determinando '
    'o seu valor em relação à outra incógnita. Em seguida, substituímos esse valor na outra equação, reduzindo o '
    'sistema a uma equação de uma variável, que pode ser resolvida facilmente.'
)

doc.add_heading('Exemplo:', level=2)
doc.add_paragraph(
    'Considere o seguinte sistema de equações:\n'
    'x + y = 12\n'
    '3x - y = 20\n\n'
    'Resolução:\n'
    '1. Começamos escolhendo a primeira equação, que é mais simples, para isolar a incógnita x. Temos:\n'
    'x = 12 - y\n\n'
    '2. Substituímos o valor de x na segunda equação:\n'
    '3(12 - y) - y = 20\n'
    '36 - 3y - y = 20\n'
    '36 - 4y = 20\n'
    '-4y = 20 - 36\n'
    '-4y = -16\n'
    'y = -16 / -4 = 4\n\n'
    '3. Agora que encontramos o valor de y, substituímos y = 4 na primeira equação:\n'
    'x + 4 = 12\n'
    'x = 12 - 4 = 8\n\n'
    'Logo, a solução para o sistema é (x, y) = (8, 4). Podemos verificar que, substituindo esses valores nas '
    'equações originais, ambas são satisfeitas.'
)

# Método da Adição
doc.add_heading('Método da Adição (ou Eliminação)', level=1)
doc.add_paragraph(
    'No Método da Adição, buscamos juntar as duas equações de modo que uma das incógnitas seja eliminada. Para '
    'isso, as equações precisam ter coeficientes opostos para uma das incógnitas. Somamos ou subtraímos as equações, '
    'eliminando uma incógnita.'
)

doc.add_heading('Exemplo:', level=2)
doc.add_paragraph(
    'Considere o mesmo sistema de equações:\n'
    'x + y = 12\n'
    '3x - y = 20\n\n'
    'Note que as incógnitas y possuem coeficientes opostos (1 e -1). Podemos somar as equações:\n'
    '(x + y) + (3x - y) = 12 + 20\n'
    'x + y + 3x - y = 32\n'
    '4x = 32\n'
    'x = 32 / 4 = 8\n\n'
    'Agora que encontramos x = 8, substituímos esse valor em uma das equações, por exemplo, na primeira:\n'
    '8 + y = 12\n'
    'y = 12 - 8 = 4\n\n'
    'A solução também é (x, y) = (8, 4), como no método da substituição.'
)

# Método da Adição com Multiplicação de Termos
doc.add_heading('Método da Adição com Multiplicação de Termos', level=1)
doc.add_paragraph(
    'Quando os coeficientes de uma incógnita não são opostos, podemos multiplicar as equações por um valor '
    'adequado para torná-los opostos.'
)

doc.add_heading('Exemplo:', level=2)
doc.add_paragraph(
    'Considere o sistema de equações:\n'
    '3x + y = 24\n'
    '5x + 2y = 60\n\n'
    'Aqui, os coeficientes de x e y não são opostos. Podemos multiplicar a primeira equação por -2 para torná-los '
    'opostos:\n'
    '-2(3x + y) = -2(24)\n'
    '-6x - 2y = -48\n\n'
    'Agora o sistema é:\n'
    '-6x - 2y = -48\n'
    '5x + 2y = 60\n\n'
    'Somando as duas equações:\n'
    '(-6x - 2y) + (5x + 2y) = -48 + 60\n'
    '-6x + 5x = 12\n'
    '-x = 12\n'
    'x = -12\n\n'
    'Agora, substituímos x = -12 em uma das equações, por exemplo, na primeira:\n'
    '3(-12) + y = 24\n'
    '-36 + y = 24\n'
    'y = 24 + 36 = 60\n\n'
    'A solução para o sistema é (x, y) = (-12, 60).'
)

# Classificação dos Sistemas de Equações
doc.add_heading('Classificação dos Sistemas de Equações', level=1)
doc.add_paragraph(
    'Um sistema de equações do 1º grau com duas incógnitas x e y, representado pelas equações a1x + b1y = c1 e a2x + b2y = c2, '
    'pode ser classificado como:\n'
    '1. Possível e Determinado: Quando o sistema tem uma única solução. A condição para isso é que as razões entre os coeficientes '
    'das incógnitas sejam diferentes:\n'
    'a1/a2 ≠ b1/b2\n\n'
    '2. Possível e Indeterminado: Quando o sistema tem infinitas soluções. A condição é que as razões entre os coeficientes '
    'das incógnitas e as constantes sejam iguais:\n'
    'a1/a2 = b1/b2 = c1/c2\n\n'
    '3. Impossível: Quando o sistema não tem solução. A condição é que as razões entre os coeficientes das incógnitas sejam iguais, '
    'mas as constantes sejam diferentes:\n'
    'a1/a2 = b1/b2 ≠ c1/c2\n'
)

# Exemplo de Classificação
doc.add_heading('Exemplo de Classificação:', level=2)
doc.add_paragraph(
    'Considere o seguinte sistema:\n'
    '2x + y = -4\n'
    '4x + 2y = 6\n\n'
    'Calculamos as razões:\n'
    'a1/a2 = 2/4 = 1/2\n'
    'b1/b2 = 1/2\n'
    'c1/c2 = -4/6 = -2/3\n\n'
    'Como a1/a2 = b1/b2, mas c1/c2 ≠ a1/a2, o sistema é **impossível**, ou seja, não possui solução.'
)

# Salvar o documento
file_path = "/mnt/data/Trabalho_Sistemas_Equacoes_1Grau_ABNT.docx"
doc.save(file_path)

file_path
